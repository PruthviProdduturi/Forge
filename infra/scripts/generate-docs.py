"""
Generate Forge platform DOCX documents.
  - Forge-Platform-Brief.docx      : One-pager with architecture diagram
  - Forge-Platform-Design-Reference.docx : Full design reference (replaces existing)

Run from repo root:
  python infra/scripts/generate_docs.py
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Colour palette
# ---------------------------------------------------------------------------
FORGE_DARK   = RGBColor(0x1a, 0x1a, 0x2e)   # dark navy
FORGE_BLUE   = RGBColor(0x0f, 0x3c, 0x78)   # deep blue
FORGE_ACCENT = RGBColor(0x00, 0x8b, 0xd4)   # azure blue
FORGE_GREEN  = RGBColor(0x10, 0x7c, 0x10)   # Microsoft green
FORGE_ORANGE = RGBColor(0xe2, 0x5a, 0x1c)   # Spark orange
FORGE_PURPLE = RGBColor(0x5c, 0x2d, 0x91)   # purple accent
FORGE_SILVER = RGBColor(0xf0, 0xf4, 0xf8)   # light background
WHITE        = RGBColor(0xff, 0xff, 0xff)
GREY_TEXT    = RGBColor(0x44, 0x44, 0x44)
LIGHT_GREY   = RGBColor(0xf8, 0xf9, 0xfa)
MED_GREY     = RGBColor(0xe9, 0xec, 0xef)
BORDER_GREY  = RGBColor(0xce, 0xd4, 0xda)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def set_cell_bg(cell, rgb: RGBColor):
    """Set cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val is not None:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_table_width(table, width_inches):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(width_inches * 1440)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


def cell_para(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.CENTER,
              italic=False, space_before=0, space_after=0):
    """Add/replace paragraph in a cell."""
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)
    p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_heading(doc, text, level=1, color=None, size=None, bold=True, space_before=12, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    elif level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = FORGE_BLUE
    return p


def add_para(doc, text, size=10, color=None, italic=False, space_before=2, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = GREY_TEXT
    return p


def add_table(doc, headers, rows, header_bg=None, col_widths=None):
    """Add a styled data table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_bg or FORGE_BLUE)
        cell_para(cell, h, bold=True, size=9, color=WHITE)

    # data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = WHITE if ri % 2 == 0 else LIGHT_GREY
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            cell_para(cell, val, size=9, color=GREY_TEXT, align=WD_ALIGN_PARAGRAPH.LEFT)

    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[ci])

    return table


def set_page_margins(doc, top=1.0, bottom=1.0, left=1.0, right=1.0):
    section = doc.sections[0]
    section.top_margin    = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin   = Inches(left)
    section.right_margin  = Inches(right)


# ---------------------------------------------------------------------------
# Architecture diagram (table-based)
# ---------------------------------------------------------------------------
def add_architecture_diagram(doc):
    """
    Renders a visual architecture diagram using nested/styled tables.

    Layout:
    ┌─────────────────────────────────────────────────────┐
    │               DEVELOPER INTERFACES                   │
    │  VS Code + Spark Connect  │  Airflow UI  │  Portal  │
    └─────────────────────────────────────────────────────┘
            │                           │
    ┌───────────────────┐   ┌───────────────────────────┐
    │  COMPUTE CLUSTER  │   │  ORCHESTRATION CLUSTER    │
    │  Spark Operator   │◄──│  Airflow                  │
    │  Spark Connect*   │   │  DQ Framework             │
    │  Trino            │   │  Developer Portal         │
    └───────────────────┘   └───────────────────────────┘
            └──────────────────────┘
                          │
    ┌─────────────────────────────────────────────────────┐
    │              ADLS GEN2 LAKEHOUSE                     │
    │    bronze/    │    silver/    │    gold/             │
    └─────────────────────────────────────────────────────┘
                          │
    ┌─────────────────────────────────────────────────────┐
    │            SHARED AZURE SERVICES                     │
    │  Key Vault │ ACR │ Log Analytics │ Grafana │ Defender│
    └─────────────────────────────────────────────────────┘
    """

    PAGE_W = 6.3  # usable width in inches

    # ── Row 0: section label ──────────────────────────────────────────────
    t0 = doc.add_table(rows=1, cols=1)
    remove_table_borders(t0)
    set_table_width(t0, PAGE_W)
    c = t0.rows[0].cells[0]
    set_cell_bg(c, FORGE_DARK)
    p = cell_para(c, 'FORGE DATA PLATFORM — ARCHITECTURE', bold=True, size=11, color=WHITE)
    c.width = Inches(PAGE_W)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Row 1: Developer Interfaces ───────────────────────────────────────
    t1 = doc.add_table(rows=2, cols=3)
    remove_table_borders(t1)
    set_table_width(t1, PAGE_W)
    # label row
    lbl = t1.rows[0].cells[0]
    t1.rows[0].cells[0].merge(t1.rows[0].cells[2])
    set_cell_bg(lbl, FORGE_BLUE)
    cell_para(lbl, 'DEVELOPER INTERFACES', bold=True, size=8, color=WHITE, space_before=3, space_after=1)

    w = PAGE_W / 3
    boxes = [
        ('VS Code + Spark Connect\n(dev — interactive)', FORGE_ACCENT),
        ('Apache Airflow UI\n(DAGs, runs, logs)', FORGE_ACCENT),
        ('Developer Portal\n(datasets, lineage, DQ, cost)', FORGE_ACCENT),
    ]
    for ci, (txt, bg) in enumerate(boxes):
        cell = t1.rows[1].cells[ci]
        set_cell_bg(cell, bg)
        set_cell_border(cell,
            top    ={'val': 'single', 'sz': 6, 'color': 'FFFFFF'},
            bottom ={'val': 'single', 'sz': 6, 'color': 'FFFFFF'},
            left   ={'val': 'single', 'sz': 6, 'color': 'FFFFFF'},
            right  ={'val': 'single', 'sz': 6, 'color': 'FFFFFF'})
        cell_para(cell, txt, size=8, color=WHITE, space_before=4, space_after=4)
        cell.width = Inches(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(1)

    # ── Row 2: Two clusters ───────────────────────────────────────────────
    t2 = doc.add_table(rows=2, cols=2)
    remove_table_borders(t2)
    set_table_width(t2, PAGE_W)

    cluster_w = PAGE_W / 2

    cluster_data = [
        {
            'label': 'COMPUTE CLUSTER',
            'label_bg': FORGE_ORANGE,
            'items': [
                '• Spark Operator (batch jobs)',
                '• Spark Connect * (dev only)',
                '• Trino (federated SQL)',
                '',
                'Node pools:',
                '  systempool  Standard_D4s_v5',
                '  sparkpool   Standard_E8s_v5  (0→5)',
                '  trinopool   Standard_D4s_v5  (0→3)',
            ],
            'body_bg': RGBColor(0xff, 0xf3, 0xeb),
        },
        {
            'label': 'ORCHESTRATION CLUSTER',
            'label_bg': FORGE_GREEN,
            'items': [
                '• Apache Airflow 3.x (KubernetesExecutor)',
                '• DQ Framework (YAML rule engine)',
                '• Developer Portal (FastAPI + Next.js)',
                '• Azure Monitor Agent (DaemonSet)',
                '',
                'Node pools:',
                '  systempool   Standard_D4s_v5',
                '  workerpool   Standard_D4s_v5  (1→3)',
            ],
            'body_bg': RGBColor(0xeb, 0xf5, 0xeb),
        },
    ]

    for ci, cd in enumerate(cluster_data):
        # label cell
        lbl_cell = t2.rows[0].cells[ci]
        set_cell_bg(lbl_cell, cd['label_bg'])
        cell_para(lbl_cell, cd['label'], bold=True, size=8, color=WHITE, space_before=3, space_after=3)
        lbl_cell.width = Inches(cluster_w)

        # body cell
        body_cell = t2.rows[1].cells[ci]
        set_cell_bg(body_cell, cd['body_bg'])
        for p in body_cell.paragraphs:
            p._element.getparent().remove(p._element)
        for line in cd['items']:
            p = body_cell.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(line)
            run.font.size = Pt(8)
            run.font.color.rgb = GREY_TEXT
        body_cell.width = Inches(cluster_w)

    doc.add_paragraph().paragraph_format.space_after = Pt(1)

    # ── Row 3: Lakehouse ──────────────────────────────────────────────────
    t3 = doc.add_table(rows=2, cols=3)
    remove_table_borders(t3)
    set_table_width(t3, PAGE_W)

    lbl3 = t3.rows[0].cells[0]
    t3.rows[0].cells[0].merge(t3.rows[0].cells[2])
    set_cell_bg(lbl3, FORGE_PURPLE)
    cell_para(lbl3, 'ADLS GEN2 LAKEHOUSE  (Delta Lake · Hierarchical Namespace)', bold=True, size=8, color=WHITE, space_before=3, space_after=1)

    zone_w = PAGE_W / 3
    zones = [
        ('bronze/',  'Raw, append-only\nImmutable source of truth\nSchema-on-read', RGBColor(0xcd, 0x7f, 0x32)),
        ('silver/',  'Cleaned, schema-enforced\nDQ-validated\nDelta Lake ACID', RGBColor(0x70, 0x7f, 0x8f)),
        ('gold/',    'Aggregated, SLA-governed\nConsumer-ready\nRead by BI / ML / Portal', RGBColor(0xc9, 0xa8, 0x00)),
    ]
    for ci, (zone, desc, bg) in enumerate(zones):
        cell = t3.rows[1].cells[ci]
        set_cell_bg(cell, bg)
        set_cell_border(cell,
            top   ={'val': 'single', 'sz': 6, 'color': 'FFFFFF'},
            bottom={'val': 'single', 'sz': 6, 'color': 'FFFFFF'},
            left  ={'val': 'single', 'sz': 6, 'color': 'FFFFFF'},
            right ={'val': 'single', 'sz': 6, 'color': 'FFFFFF'})
        for p in cell.paragraphs:
            p._element.getparent().remove(p._element)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(zone)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = WHITE
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after = Pt(3)
        r2 = p2.add_run(desc)
        r2.font.size = Pt(7)
        r2.font.color.rgb = WHITE
        cell.width = Inches(zone_w)

    doc.add_paragraph().paragraph_format.space_after = Pt(1)

    # ── Row 4: Shared Azure Services ──────────────────────────────────────
    t4 = doc.add_table(rows=2, cols=5)
    remove_table_borders(t4)
    set_table_width(t4, PAGE_W)

    lbl4 = t4.rows[0].cells[0]
    t4.rows[0].cells[0].merge(t4.rows[0].cells[4])
    set_cell_bg(lbl4, FORGE_DARK)
    cell_para(lbl4, 'SHARED AZURE SERVICES', bold=True, size=8, color=WHITE, space_before=3, space_after=1)

    svc_w = PAGE_W / 5
    services = [
        ('Key Vault', 'Secrets +\nHSM keys', RGBColor(0x00, 0x78, 0xd4)),
        ('ACR', 'Private\nimage registry', RGBColor(0x00, 0x78, 0xd4)),
        ('Log Analytics\n+ Grafana', 'Metrics, logs\nalerts, dashboards', RGBColor(0x00, 0x78, 0xd4)),
        ('Defender', 'Runtime threat\ndetection (S360)', RGBColor(0x00, 0x78, 0xd4)),
        ('Private DNS\n+ Endpoints', 'No public\ndata plane', RGBColor(0x00, 0x78, 0xd4)),
    ]
    for ci, (name, desc, bg) in enumerate(services):
        cell = t4.rows[1].cells[ci]
        set_cell_bg(cell, bg)
        set_cell_border(cell,
            top   ={'val': 'single', 'sz': 4, 'color': 'FFFFFF'},
            bottom={'val': 'single', 'sz': 4, 'color': 'FFFFFF'},
            left  ={'val': 'single', 'sz': 4, 'color': 'FFFFFF'},
            right ={'val': 'single', 'sz': 4, 'color': 'FFFFFF'})
        for p in cell.paragraphs:
            p._element.getparent().remove(p._element)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(name)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = WHITE
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(3)
        r2 = p2.add_run(desc)
        r2.font.size = Pt(7)
        r2.font.color.rgb = RGBColor(0xcc, 0xe4, 0xff)
        cell.width = Inches(svc_w)

    # footnote
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    fn = doc.add_paragraph()
    fn.paragraph_format.space_before = Pt(0)
    fn.paragraph_format.space_after = Pt(0)
    r = fn.add_run('* Spark Connect available in dev environment only — interactive VS Code sessions against live cluster.')
    r.font.size = Pt(7)
    r.italic = True
    r.font.color.rgb = GREY_TEXT


# ---------------------------------------------------------------------------
# Document 1 — Platform Brief (one-pager)
# ---------------------------------------------------------------------------
def build_platform_brief(output_path):
    doc = Document()
    set_page_margins(doc, top=0.7, bottom=0.7, left=0.9, right=0.9)

    # ── Title block ───────────────────────────────────────────────────────
    title_tbl = doc.add_table(rows=1, cols=1)
    remove_table_borders(title_tbl)
    set_table_width(title_tbl, 6.3)
    tc = title_tbl.rows[0].cells[0]
    set_cell_bg(tc, FORGE_DARK)
    for p in tc.paragraphs:
        p._element.getparent().remove(p._element)
    p = tc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('FORGE')
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = FORGE_ACCENT
    p2 = tc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run('Core Platform  ·  Platform Brief')
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0xcc, 0xdd, 0xee)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── What is Forge ─────────────────────────────────────────────────────
    add_para(doc,
        'Forge is the core platform that moves data from raw source systems to '
        'governed, serving-ready analytics — reliably, observably, and at scale. It is the single '
        'system through which all data enters the lakehouse, gets validated, and becomes trusted '
        'for downstream consumption.',
        size=10, space_before=0, space_after=6)

    # ── Architecture diagram ──────────────────────────────────────────────
    add_architecture_diagram(doc)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Two-column section: Problems Solved + Security ────────────────────
    t_cols = doc.add_table(rows=1, cols=2)
    remove_table_borders(t_cols)
    set_table_width(t_cols, 6.3)
    left_cell  = t_cols.rows[0].cells[0]
    right_cell = t_cols.rows[0].cells[1]
    left_cell.width  = Inches(3.1)
    right_cell.width = Inches(3.1)

    # left: problems solved
    for p in left_cell.paragraphs:
        p._element.getparent().remove(p._element)

    def add_cell_heading(cell, text):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = FORGE_BLUE

    add_cell_heading(left_cell, 'WHAT IT SOLVES')
    problems = [
        ('Unreliable source data',          'DQ gates at every layer — nothing untested reaches gold'),
        ('No production job governance',    'Spark Operator: every job scheduled, tracked, auditable'),
        ('Unknown data lineage',            'OpenLineage on every job — column-level lineage to gold'),
        ('Secrets in code',                 'Workload Identity + Key Vault — zero long-lived credentials'),
        ('Slow developer setup',            'VS Code + Spark Connect — live cluster in minutes'),
        ('Silent pipeline failures',        'Azure Monitor Alerts fire before consumers notice'),
    ]
    for prob, sol in problems:
        p = left_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f'• {prob}')
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = FORGE_DARK
        p2 = left_cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(f'  {sol}')
        r2.font.size = Pt(8)
        r2.italic = True
        r2.font.color.rgb = GREY_TEXT

    # right: security + infra
    for p in right_cell.paragraphs:
        p._element.getparent().remove(p._element)

    add_cell_heading(right_cell, 'SECURITY (S360)')
    security = [
        ('No long-lived credentials',   'OIDC Workload Identity — pods exchange K8s tokens for Azure tokens'),
        ('No secrets in code',          'Key Vault + CSI driver — secrets as in-memory pod volumes'),
        ('Least-privilege identities',  '5 workload identities, each scoped to minimum required containers'),
        ('No public data plane',        'ADLS, KV, ACR all behind private endpoints'),
        ('Full audit trail',            'Every kubectl, secret access, and storage op logged in Log Analytics'),
        ('Threat detection',            'Defender for Containers, Storage, Key Vault at subscription scope'),
    ]
    for item, detail in security:
        p = right_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f'• {item}')
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = FORGE_DARK
        p2 = right_cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(f'  {detail}')
        r2.font.size = Pt(8)
        r2.italic = True
        r2.font.color.rgb = GREY_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Developer workflow ────────────────────────────────────────────────
    add_heading(doc, 'DEVELOPER WORKFLOW', level=3, size=9, space_before=4, space_after=3)
    steps = [
        ('1  Write',   'VS Code notebook → Spark Connect → live dev cluster (real data, real results)'),
        ('2  Test',    'forge generate scaffold → add business logic → run DQ locally against dev cluster'),
        ('3  Review',  'PR → CI: lint, unit tests, DQ validation, image scan, IaC policy check'),
        ('4  Deploy',  'Merge to main → auto-deploy to dev  (DAG changes live within 30 seconds)'),
        ('5  Promote', 'Release tag → manual approval gate → production deployment'),
    ]
    wf_tbl = doc.add_table(rows=len(steps), cols=2)
    remove_table_borders(wf_tbl)
    set_table_width(wf_tbl, 6.3)
    for ri, (step, desc) in enumerate(steps):
        bg = FORGE_ACCENT if ri % 2 == 0 else FORGE_BLUE
        sc = wf_tbl.rows[ri].cells[0]
        dc = wf_tbl.rows[ri].cells[1]
        sc.width = Inches(0.9)
        dc.width = Inches(5.4)
        set_cell_bg(sc, bg)
        set_cell_bg(dc, LIGHT_GREY if ri % 2 == 0 else WHITE)
        set_cell_border(sc, top={'val':'single','sz':4,'color':'FFFFFF'}, bottom={'val':'single','sz':4,'color':'FFFFFF'})
        set_cell_border(dc, top={'val':'single','sz':4,'color':'FFFFFF'}, bottom={'val':'single','sz':4,'color':'FFFFFF'})
        cell_para(sc, step, bold=True, size=8, color=WHITE, space_before=3, space_after=3)
        cell_para(dc, desc, size=8, color=GREY_TEXT, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=3, space_after=3)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Footer / where to start ───────────────────────────────────────────
    footer_tbl = doc.add_table(rows=1, cols=1)
    remove_table_borders(footer_tbl)
    set_table_width(footer_tbl, 6.3)
    fc = footer_tbl.rows[0].cells[0]
    set_cell_bg(fc, MED_GREY)
    for p in fc.paragraphs:
        p._element.getparent().remove(p._element)
    p = fc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(
        'Infrastructure: Bicep (Azure)  ·  Kubernetes: AKS 1.32  ·  Storage: ADLS Gen2 + Delta Lake 4.0  ·  '
        'Environments: dev (alias-scoped) → prod (release-gated)  ·  Compliance: S360'
    )
    r.font.size = Pt(8)
    r.font.color.rgb = GREY_TEXT

    doc.save(output_path)
    print(f'  ✓  {output_path}')


# ---------------------------------------------------------------------------
# Document 2 — Platform Design Reference
# ---------------------------------------------------------------------------
def build_design_reference(output_path):
    doc = Document()
    set_page_margins(doc, top=1.0, bottom=1.0, left=1.1, right=1.1)

    # Title
    title_tbl = doc.add_table(rows=1, cols=1)
    remove_table_borders(title_tbl)
    set_table_width(title_tbl, 6.3)
    tc = title_tbl.rows[0].cells[0]
    set_cell_bg(tc, FORGE_DARK)
    for p in tc.paragraphs:
        p._element.getparent().remove(p._element)
    for text, sz, clr in [
        ('FORGE', 24, FORGE_ACCENT),
        ('Platform Design Reference', 13, RGBColor(0xcc, 0xdd, 0xee)),
        ('Version 1.0  ·  Internal — Platform Team', 9, RGBColor(0x88, 0xaa, 0xcc)),
    ]:
        p = tc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4 if sz < 20 else 10)
        p.paragraph_format.space_after  = Pt(4 if sz < 20 else 2)
        r = p.add_run(text)
        r.bold = sz >= 20
        r.font.size = Pt(sz)
        r.font.color.rgb = clr

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ── 1. Executive Summary ──────────────────────────────────────────────
    add_heading(doc, '1.  Executive Summary', level=1)
    add_para(doc,
        'Forge is the core data engineering platform that moves data from raw sources to governed, '
        'serving-ready analytics. It is built around three principles: separate compute from '
        'orchestration, use Azure-native observability, and put everything in Git.',
        size=10)
    add_para(doc,
        'Two AKS clusters per environment provide isolation between elastic data processing '
        '(Spark, Trino) and stable control-plane services (Airflow, DQ, Portal). Failures and '
        'scaling events in one cluster do not affect the other.',
        size=10)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── 2. Architecture Principles ────────────────────────────────────────
    add_heading(doc, '2.  Architecture Principles', level=1)
    add_table(doc,
        headers=['Principle', 'What it means in practice'],
        rows=[
            ('Separation of concerns',      'Orchestration and compute are isolated clusters. A Spark OOM does not page the Airflow on-call.'),
            ('Minimal blast radius',        'Each cluster, layer, and job runs with minimum permissions. No shared credentials.'),
            ('Environment isolation',       'dev and prod are independent deployments — data, identities, and secrets are never shared.'),
            ('Everything as code',          'Infrastructure is Bicep. Platform config is Helm. Pipelines are Python DAGs. Quality rules are YAML.'),
            ('Observable by default',       'Every pipeline run, query, and data write is observable without SSH or tribal knowledge.'),
            ('Developer experience first',  'Engineers write and test Spark code from VS Code against a live dev cluster. No local Spark.'),
            ('Secure by default',           'Workload Identity replaces secrets. Private endpoints on every PaaS. Zero long-lived credentials.'),
        ],
        col_widths=[2.0, 4.3],
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── 3. Reference Architecture ─────────────────────────────────────────
    add_heading(doc, '3.  Reference Architecture', level=1)
    add_architecture_diagram(doc)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── 4. Two-Cluster Model ──────────────────────────────────────────────
    add_heading(doc, '4.  Two-Cluster Model', level=1)
    add_table(doc,
        headers=['Cluster', 'Name', 'Purpose', 'Key Components'],
        rows=[
            ('Compute',        'aks-forge-compute-{alias}-{env}',        'Elastic data processing',     'Spark Operator, Spark Connect (dev), Trino'),
            ('Orchestration',  'aks-forge-orchestration-{alias}-{env}',  'Stable control-plane services', 'Airflow, DQ runner, Developer Portal'),
        ],
        col_widths=[1.2, 2.3, 1.6, 1.5],
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_heading(doc, 'Node Pools — Compute', level=2, size=11)
    add_table(doc,
        headers=['Pool', 'VM Size', 'Min', 'Max (dev)', 'Workload'],
        rows=[
            ('systempool', 'Standard_D4s_v5', '1', '2',  'System addons only (CriticalAddonsOnly taint)'),
            ('sparkpool',  'Standard_E8s_v5', '0', '5',  'Spark executors — memory-optimised'),
            ('trinopool',  'Standard_D4s_v5', '0', '3',  'Trino workers'),
        ],
        col_widths=[1.3, 1.8, 0.5, 0.9, 2.0],
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_heading(doc, 'Node Pools — Orchestration', level=2, size=11)
    add_table(doc,
        headers=['Pool', 'VM Size', 'Min', 'Max (dev)', 'Workload'],
        rows=[
            ('systempool',  'Standard_D4s_v5', '1', '2', 'System addons only'),
            ('workerpool',  'Standard_D4s_v5', '1', '3', 'Airflow, DQ runner, Developer Portal'),
        ],
        col_widths=[1.3, 1.8, 0.5, 0.9, 2.0],
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── 5. Medallion Lakehouse ────────────────────────────────────────────
    add_heading(doc, '5.  Medallion Lakehouse (ADLS Gen2)', level=1)
    add_table(doc,
        headers=['Layer', 'Container', 'Format', 'Purpose', 'Writers', 'Readers'],
        rows=[
            ('Bronze', 'bronze/',       'Delta / raw',  'Immutable source of truth, append-only',   'Airflow',       'Spark'),
            ('Silver', 'silver/',       'Delta Lake',   'Cleaned, schema-enforced, DQ-validated',    'Spark',         'Spark, Trino'),
            ('Gold',   'gold/',         'Delta Lake',   'Aggregated, SLA-governed, consumer-ready',  'Spark',         'Trino, Portal'),
            ('Code',   'code/',         'Files',        'Job notebooks, JARs, runner scripts, checkpoints/<pipeline_id>/', 'CI/CD, Spark', 'Spark'),
        ],
        col_widths=[0.7, 0.9, 0.9, 2.2, 0.8, 0.8],
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── 6. Managed Identities ─────────────────────────────────────────────
    add_heading(doc, '6.  Managed Identities', level=1)
    add_para(doc, '9 identities per environment — 4 AKS infrastructure, 5 workload.', size=10)
    add_table(doc,
        headers=['Identity', 'Type', 'Storage Access', 'KV'],
        rows=[
            ('id-aks-controlplane-compute-{alias}-{env}',       'AKS infra',  'None',                                          '—'),
            ('id-aks-kubelet-compute-{alias}-{env}',            'AKS infra',  'None (ACR pull via separate assignment)',        '—'),
            ('id-aks-controlplane-orchestration-{alias}-{env}', 'AKS infra',  'None',                                          '—'),
            ('id-aks-kubelet-orchestration-{alias}-{env}',      'AKS infra',  'None (ACR pull)',                               '—'),
            ('id-forge-spark-{alias}-{env}',                    'Workload',   'Contributor: bronze, silver, gold, code',        'Secrets User'),
            ('id-forge-trino-{alias}-{env}',                    'Workload',   'Reader: silver, gold',                          'Secrets User'),
            ('id-forge-airflow-{alias}-{env}',                  'Workload',   'Contributor: bronze · Reader: code',            'Secrets User'),
            ('id-forge-dq-{alias}-{env}',                       'Workload',   'Reader: bronze, silver, gold',                  'Secrets User'),
            ('id-forge-portal-{alias}-{env}',                   'Workload',   'Reader: gold',                                  'Secrets User'),
        ],
        col_widths=[2.8, 0.8, 2.2, 1.0],
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── 7. Security (S360) ────────────────────────────────────────────────
    add_heading(doc, '7.  Security & S360 Compliance', level=1)
    add_table(doc,
        headers=['S360 Control', 'Status', 'Implementation'],
        rows=[
            ('No long-lived credentials',        '✓ Compliant', 'Azure Workload Identity (OIDC federation)'),
            ('No public data plane endpoints',   '✓ Compliant', 'Private endpoints on ADLS, KV, ACR; AKS API server public but AAD-gated'),
            ('Secrets in managed vault',         '✓ Compliant', 'Azure Key Vault Premium + CSI driver'),
            ('Encryption at rest',               '✓ Compliant', 'AES-256 on all services'),
            ('Encryption in transit',            '✓ Compliant', 'TLS 1.2+ enforced everywhere'),
            ('Vulnerability scanning',           '✓ Compliant', 'ACR Defender + OPA Gatekeeper admission control'),
            ('Audit logging',                    '✓ Compliant', 'Log Analytics, 90 days minimum retention'),
            ('Threat detection',                 '✓ Compliant', 'Defender for Containers, Storage, Key Vault, CSPM'),
            ('RBAC via Azure AD',                '✓ Compliant', 'All roles backed by AAD; disableLocalAccounts on AKS'),
            ('ipTag on public IPs (NS2.1.1)',     '✓ Compliant', 'Pre-created PIPs with FirstPartyUsage=/NonProd|Prod'),
        ],
        header_bg=FORGE_GREEN,
        col_widths=[2.5, 1.0, 3.0],
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── 8. Observability ──────────────────────────────────────────────────
    add_heading(doc, '8.  Observability', level=1)
    add_table(doc,
        headers=['Workspace', 'What feeds into it'],
        rows=[
            ('law-forge-compute-{alias}-{env}',        'Compute AKS: apiserver, audit, scheduler, controller logs; Defender; OMS metrics'),
            ('law-forge-orchestration-{alias}-{env}',  'Orchestration AKS logs; ADLS diagnostics; Key Vault audit logs'),
            ('law-forge-platform-{alias}-{env}',       'VNet flow data, NSG diagnostics'),
        ],
        col_widths=[2.8, 3.5],
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── 9. Delivery Lifecycle ─────────────────────────────────────────────
    add_heading(doc, '9.  Delivery Lifecycle', level=1)
    add_table(doc,
        headers=['Change type', 'Deploy mechanism', 'Speed'],
        rows=[
            ('DAG-only',      'Airflow git-sync',                            '30 seconds'),
            ('Spark job',     'CD pipeline uploads notebook to ADLS code/',  '~5 minutes'),
            ('Helm config',   'CD pipeline helm upgrade',                    '~10 minutes'),
            ('Bicep infra',   'CD pipeline az deployment sub create',        '~20 minutes'),
            ('Image rebuild', 'CD: build → scan → push → helm upgrade',      '~30 minutes'),
        ],
        col_widths=[1.5, 3.3, 1.5],
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── 10. Document Index ────────────────────────────────────────────────
    add_heading(doc, '10.  Architecture Document Index', level=1)
    add_table(doc,
        headers=['#', 'Document', 'What it covers'],
        rows=[
            ('1',  'docs/architecture/01-overview.md',            'Infrastructure reference — resource groups, clusters, networking, identities'),
            ('2',  'docs/architecture/02-rg-inventory.md',        'Every Azure resource, which RG, why it exists, what uses it'),
            ('3',  'docs/architecture/03-networking.md',          'VNet layout, private endpoints, DNS, Calico policies, NSGs'),
            ('4',  'docs/architecture/04-security-s360.md',       'S360 control mapping — identity, network, secrets, encryption, scanning'),
            ('5',  'docs/architecture/05-storage.md',             'Medallion zones, partitioning standard, run trackers, versioning'),
            ('6',  'docs/architecture/06-compute.md',             'Spark Operator, Spark Connect, Trino, node pools, ADLS access'),
            ('7',  'docs/architecture/07-orchestration.md',       'Airflow KubernetesExecutor, DAG git-sync, Key Vault secrets backend'),
            ('8',  'docs/architecture/08-observability.md',       'Azure Monitor, Grafana, Log Analytics, SLOs, cost telemetry'),
            ('9',  'docs/architecture/09-dq-framework.md',        'DQ rule types, YAML format, severity gating, results store'),
            ('10', 'docs/architecture/10-lineage.md',             'OpenLineage, Microsoft Purview, column-level lineage, impact analysis'),
            ('11', 'docs/architecture/11-developer-portal.md',    'Portal API + Web, auth flow, all 6 API domains'),
            ('12', 'docs/architecture/12-end-to-end-flow.md',     'Full data flow: source systems → bronze → silver → gold'),
            ('13', 'docs/architecture/13-restatement.md',         'Partition restatement, backfill, Restatement Registry, safety guards'),
            ('14', 'docs/architecture/14-environment-promotion.md','Dev vs prod, Spark Connect vs Operator, PR→CI→dev→prod promotion'),
        ],
        col_widths=[0.3, 3.2, 3.0],
    )

    doc.save(output_path)
    print(f'  ✓  {output_path}')


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
if __name__ == '__main__':
    repo_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    ref_dir = os.path.join(repo_root, 'docs', 'reference')
    os.makedirs(ref_dir, exist_ok=True)

    print('Generating Forge DOCX documents...')
    build_platform_brief(os.path.join(ref_dir, 'Forge-Platform-Brief.docx'))
    build_design_reference(os.path.join(ref_dir, 'Forge-Platform-Design-Reference.docx'))
    print('Done.')
