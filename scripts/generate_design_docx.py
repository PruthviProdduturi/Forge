"""
Generate Forge-Platform-Design-Reference.docx from docs/DESIGN.md
Run from repo root: python scripts/generate_design_docx.py
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x0A, 0x2F, 0x5A)
ACCENT = RGBColor(0x00, 0x78, 0xD4)
GREY = RGBColor(0x64, 0x74, 0x8B)
CODE_BG = RGBColor(0xF1, 0xF5, 0xF9)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_cover(doc: Document, title: str, subtitle: str):
    doc.add_section()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)

    run = p.add_run("⬡")
    run.font.size = Pt(48)
    run.font.color.rgb = ACCENT

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(24)
    r2 = p2.add_run(title)
    r2.font.size = Pt(28)
    r2.font.bold = True
    r2.font.color.rgb = NAVY

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(subtitle)
    r3.font.size = Pt(13)
    r3.font.color.rgb = GREY

    meta = [
        ("Version", "1.0"),
        ("Status", "Production"),
        ("Classification", "Internal — Platform Team"),
        ("Last Updated", "2026-03-25"),
    ]
    for label, value in meta:
        pm = doc.add_paragraph()
        pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pm.paragraph_format.space_before = Pt(4)
        rm = pm.add_run(f"{label}:  ")
        rm.font.bold = True
        rm.font.size = Pt(11)
        rm.font.color.rgb = NAVY
        rv = pm.add_run(value)
        rv.font.size = Pt(11)
        rv.font.color.rgb = GREY

    doc.add_page_break()


def style_heading(para, level: int):
    for run in para.runs:
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(20)
            run.font.color.rgb = NAVY
        elif level == 2:
            run.font.size = Pt(15)
            run.font.color.rgb = NAVY
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = ACCENT


def add_code_block(doc: Document, text: str):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.right_indent = Inches(0.25)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    # Shade background via XML
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F1F5F9")
    pPr.append(shd)

    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)


def render_inline(text: str, run):
    """Apply bold/italic/code to a run — simplified: strip markers, set bold."""
    # Strip bold markers for rendering
    cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    cleaned = re.sub(r"`(.+?)`", r"\1", cleaned)
    cleaned = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", cleaned)
    run.text = cleaned


def add_table_from_md(doc: Document, lines: list[str]):
    # Find header, separator, rows
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith("|") and not re.match(r"^\|[-| :]+\|$", line):
            cells = [c.strip() for c in line.strip("|").split("|")]
            rows.append(cells)

    if not rows:
        return

    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            # Clean markdown from cell text
            cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text)
            cleaned = re.sub(r"`(.+?)`", r"\1", cleaned)
            cleaned = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", cleaned)
            cell.text = cleaned
            if r_idx == 0:
                set_cell_bg(cell, NAVY)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = WHITE
                        run.font.size = Pt(9)
            else:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)


def generate(md_path: Path, out_path: Path):
    content = md_path.read_text(encoding="utf-8")
    lines = content.splitlines()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_cover(
        doc,
        "Forge — Platform Design Reference",
        "Core Data Engineering Platform",
    )

    i = 0
    in_code = False
    code_lines = []
    table_lines = []
    in_table = False

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                add_code_block(doc, "\n".join(code_lines))
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Table detection
        if line.startswith("|"):
            table_lines.append(line)
            i += 1
            continue
        else:
            if table_lines:
                add_table_from_md(doc, table_lines)
                table_lines = []

        # Skip separators and TOC
        if re.match(r"^---+$", line.strip()):
            i += 1
            continue

        # Skip frontmatter blockquote lines at top
        if line.startswith("> "):
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # Skip table of contents heading and entries
            if text == "Table of Contents":
                # Skip until next blank line
                while i < len(lines) and lines[i].strip() != "---":
                    i += 1
                continue
            if level == 1:
                para = doc.add_heading(text, level=1)
            elif level == 2:
                para = doc.add_heading(text, level=2)
            elif level == 3:
                para = doc.add_heading(text, level=3)
            else:
                para = doc.add_heading(text, level=4)
            style_heading(para, level)
            i += 1
            continue

        # Bullet points
        m = re.match(r"^[-*]\s+(.*)", line)
        if m:
            text = m.group(1)
            cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            cleaned = re.sub(r"`(.+?)`", r"\1", cleaned)
            cleaned = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", cleaned)
            para = doc.add_paragraph(cleaned, style="List Bullet")
            para.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Numbered lists
        m = re.match(r"^\d+\.\s+(.*)", line)
        if m:
            text = m.group(1)
            cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            cleaned = re.sub(r"`(.+?)`", r"\1", cleaned)
            cleaned = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", cleaned)
            para = doc.add_paragraph(cleaned, style="List Number")
            para.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Arrow lines (→ links)
        if line.startswith("→"):
            cleaned = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line)
            para = doc.add_paragraph(cleaned)
            for run in para.runs:
                run.font.color.rgb = ACCENT
                run.font.italic = True
                run.font.size = Pt(10)
            i += 1
            continue

        # Blank line
        if line.strip() == "":
            i += 1
            continue

        # Normal paragraph
        cleaned = re.sub(r"\*\*(.+?)\*\*", lambda m: m.group(1), line)
        cleaned = re.sub(r"`(.+?)`", lambda m: m.group(1), cleaned)
        cleaned = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", cleaned)
        if cleaned.strip():
            para = doc.add_paragraph(cleaned)
            para.paragraph_format.space_after = Pt(4)
        i += 1

    # Flush any remaining table
    if table_lines:
        add_table_from_md(doc, table_lines)

    doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    repo_root = Path(__file__).parent.parent
    md_path = repo_root / "docs" / "DESIGN.md"
    out_path = repo_root / "docs" / "Forge-Platform-Design-Reference.docx"
    generate(md_path, out_path)
